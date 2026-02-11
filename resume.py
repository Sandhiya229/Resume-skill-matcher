import streamlit as st
import pdfplumber
import docx
from docx import Document
from io import BytesIO
import matplotlib.pyplot as plt
import re


# ---------------- PAGE SETUP ----------------

st.set_page_config(
    page_title="Resume Skill Matcher",
    page_icon="📄",
    layout="wide"
)

st.title("📄 Resume Skill Matcher & Career Guide")


# ---------------- ROLE DATABASE ----------------

ROLE_DATA = {

    "Python Developer": {
        "skills": ["python", "django", "flask", "api", "mysql", "git", "oop"],

        "roadmap": [
            "Learn Python Basics",
            "Master OOP",
            "Learn Django / Flask",
            "Build REST APIs",
            "Work with Databases",
            "Deploy Projects"
        ],

        "projects": [
            "Student Management System",
            "REST API using Flask",
            "Django Blog App",
            "Automation Scripts",
            "Weather App using API"
        ]
    },

    "Frontend Developer": {
        "skills": ["html", "css", "javascript", "react", "bootstrap", "git"],

        "roadmap": [
            "HTML & CSS",
            "JavaScript",
            "React",
            "UI Frameworks",
            "Git & GitHub",
            "Build Projects"
        ],

        "projects": [
            "Portfolio Website",
            "React Dashboard",
            "E-Commerce UI",
            "Landing Page Clone",
            "Todo App"
        ]
    },

    "Data Analyst": {
        "skills": ["python", "sql", "excel", "power bi", "tableau", "statistics"],

        "roadmap": [
            "Excel & SQL",
            "Python Analysis",
            "Data Cleaning",
            "Visualization",
            "Power BI / Tableau",
            "Case Studies"
        ],

        "projects": [
            "Sales Dashboard",
            "Customer Segmentation",
            "Stock Analysis",
            "Covid Data Analysis",
            "Financial Report System"
        ]
    }
}


# ---------------- PROFESSIONAL SAMPLE RESUMES ----------------

SAMPLE_RESUME = {

    # ---------- PYTHON ----------

    "Python Developer": {

        "name": "Johnathan Doe",
        "title": "Python Backend Developer",

        "contact": "📍 Chennai, India | 📧 john@email.com | 📱 9876543210 | GitHub | LinkedIn",

        "summary": (
            "Highly motivated Python Developer with 3+ years of experience in "
            "backend development, REST API design, and database management. "
            "Skilled in building scalable web applications using Django and Flask."
        ),

        "skills": [
            "Programming: Python, OOP, Data Structures",
            "Frameworks: Django, Flask, FastAPI",
            "Databases: MySQL, PostgreSQL, MongoDB",
            "APIs: RESTful API Development, JSON",
            "Tools: Git, GitHub, Docker, VS Code",
            "Cloud: AWS Basics, CI/CD"
        ],

        "experience": [

            "Senior Python Developer – ABC Tech (2022–Present)\n"
            "• Built scalable APIs for 50K+ users\n"
            "• Improved performance by 35%\n"
            "• Led development team\n"
            "• Integrated payment gateways",

            "Python Developer – XYZ Solutions (2020–2022)\n"
            "• Developed Django modules\n"
            "• Automated workflows\n"
            "• Optimized databases\n"
            "• Reduced bugs by 40%"
        ],

        "projects": [
            "Resume Analyzer System",
            "E-Commerce Backend",
            "Student Management System",
            "Weather App"
        ],

        "achievements": [
            "Employee of the Year – 2023",
            "Top Performer Award",
            "500+ LeetCode Problems"
        ],

        "education": "B.Sc Computer Science – University of Madras – 2020",

        "certifications": [
            "Google Python Certificate",
            "Advanced Django – Udemy",
            "REST API Security – Coursera"
        ]
    },


    # ---------- FRONTEND ----------

    "Frontend Developer": {

        "name": "Sophia Williams",
        "title": "Frontend Engineer",

        "contact": "📍 Bangalore, India | 📧 sophia@email.com | 📱 9876543211 | GitHub | LinkedIn",

        "summary": (
            "Creative Frontend Developer with 3+ years of experience in building "
            "responsive and interactive web applications using React and JavaScript."
        ),

        "skills": [
            "Web: HTML5, CSS3, JavaScript ES6",
            "Frameworks: React, Bootstrap, Tailwind",
            "UI/UX: Responsive Design, Figma",
            "Tools: Git, GitHub, VS Code",
            "Testing: Jest, Lighthouse",
            "Performance Optimization"
        ],

        "experience": [

            "Frontend Engineer – ABC Web (2021–Present)\n"
            "• Built React dashboards\n"
            "• Improved UX by 40%\n"
            "• Integrated APIs\n"
            "• Optimized page load",

            "UI Developer – XYZ Digital (2019–2021)\n"
            "• Designed landing pages\n"
            "• Fixed UI bugs\n"
            "• Improved accessibility\n"
            "• Worked with designers"
        ],

        "projects": [
            "Portfolio Website",
            "React Admin Dashboard",
            "E-Commerce UI",
            "Social Media UI Clone"
        ],

        "achievements": [
            "Best UI Design Award – 2022",
            "Top Performer – ABC Web",
            "100+ UI Components Built"
        ],

        "education": "B.Sc Information Technology – Anna University – 2019",

        "certifications": [
            "Frontend Masters – React",
            "Google UX Certificate",
            "FreeCodeCamp Responsive Design"
        ]
    },


    # ---------- DATA ANALYST ----------

    "Data Analyst": {

        "name": "Michael Anderson",
        "title": "Senior Data Analyst",

        "contact": "📍 Hyderabad, India | 📧 michael@email.com | 📱 9876543222 | GitHub | LinkedIn",

        "summary": (
            "Detail-oriented Data Analyst with strong expertise in data visualization, "
            "statistical analysis, and business intelligence reporting."
        ),

        "skills": [
            "Analysis: Python, Pandas, NumPy",
            "Databases: SQL, MySQL, PostgreSQL",
            "Visualization: Power BI, Tableau",
            "Statistics: Hypothesis Testing, Regression",
            "Tools: Excel, Jupyter, PowerPoint",
            "ETL & Data Cleaning"
        ],

        "experience": [

            "Senior Data Analyst – ABC Analytics (2020–Present)\n"
            "• Created dashboards for management\n"
            "• Improved decision making\n"
            "• Built forecasting models\n"
            "• Led analytics team",

            "Junior Analyst – XYZ Corp (2018–2020)\n"
            "• Cleaned large datasets\n"
            "• Generated reports\n"
            "• Automated Excel tasks\n"
            "• Reduced manual work"
        ],

        "projects": [
            "Sales Performance Dashboard",
            "Customer Segmentation",
            "Stock Market Analysis",
            "Covid Impact Study"
        ],

        "achievements": [
            "Best Analyst Award – 2021",
            "Top Performer – XYZ Corp",
            "Handled 10M+ Records"
        ],

        "education": "B.Sc Statistics – Osmania University – 2018",

        "certifications": [
            "Google Data Analytics",
            "Microsoft Power BI",
            "SQL for Data Science"
        ]
    }
}


# ---------------- FUNCTIONS ----------------


def read_pdf(file):

    text = ""

    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""

    return text.lower()


def read_docx(file):

    doc = docx.Document(file)

    text = ""

    for p in doc.paragraphs:
        text += p.text + " "

    return text.lower()


def clean_text(text):

    text = re.sub(r'[^a-zA-Z ]', ' ', text)

    return text.lower()


def analyze_resume(text, role):

    skills = ROLE_DATA[role]["skills"]

    matched = []
    missing = []

    for s in skills:

        if s in text:
            matched.append(s)

        else:
            missing.append(s)

    percent = int((len(matched) / len(skills)) * 100)

    return percent, matched, missing


def generate_docx(data):

    doc = Document()

    doc.add_heading(data["name"], level=1)

    doc.add_paragraph(data["title"])
    doc.add_paragraph(data["contact"])


    doc.add_heading("Professional Summary", level=2)
    doc.add_paragraph(data["summary"])


    doc.add_heading("Technical Skills", level=2)
    for s in data["skills"]:
        doc.add_paragraph(f"• {s}")


    doc.add_heading("Professional Experience", level=2)
    for e in data["experience"]:
        doc.add_paragraph(e)


    doc.add_heading("Key Projects", level=2)
    for p in data["projects"]:
        doc.add_paragraph(f"• {p}")


    doc.add_heading("Achievements", level=2)
    for a in data["achievements"]:
        doc.add_paragraph(f"• {a}")


    doc.add_heading("Education", level=2)
    doc.add_paragraph(data["education"])


    doc.add_heading("Certifications", level=2)
    for c in data["certifications"]:
        doc.add_paragraph(f"• {c}")


    return doc


# ---------------- UI ----------------


role = st.selectbox("🎯 Select Job Role", ROLE_DATA.keys())

upload = st.file_uploader(
    "📤 Upload Resume (PDF / DOCX)",
    type=["pdf", "docx"]
)


# ---------------- ANALYZE ----------------


if st.button("🚀 Analyze Resume"):

    if upload is None:
        st.warning("Upload resume first")
        st.stop()


    if upload.type == "application/pdf":
        text = read_pdf(upload)

    else:
        text = read_docx(upload)


    text = clean_text(text)

    score, matched, missing = analyze_resume(text, role)


    # ---------- RESULT ----------

    st.subheader("📊 Resume Result")

    st.metric("Match Percentage", f"{score}%")


    col1, col2 = st.columns(2)


    with col1:

        st.success("Matched Skills")

        for s in matched:
            st.write("✅", s)


    with col2:

        st.error("Missing Skills")

        for s in missing:
            st.write("❌", s)


    # ---------- CHART ----------

    st.subheader("📈 Skill Chart")

    labels = ["Matched", "Missing"]
    sizes = [len(matched), len(missing)]

    fig, ax = plt.subplots()

    ax.pie(sizes, labels=labels, autopct="%1.1f%%")

    ax.axis("equal")

    st.pyplot(fig)


    # ---------- ROADMAP ----------

    st.subheader("🛣️ Career Roadmap")

    for i, step in enumerate(ROLE_DATA[role]["roadmap"], 1):

        st.write(f"Step {i} → {step}")


    # ---------- PROJECT SUGGESTIONS ----------

    st.subheader("💡 Suggested Projects")

    for p in ROLE_DATA[role]["projects"]:

        st.write("📌", p)


    # ---------- DOWNLOAD SAMPLE RESUME ----------

    st.subheader("📥 Example Resume")

    sample = SAMPLE_RESUME[role]

    doc = generate_docx(sample)

    file = BytesIO()

    doc.save(file)

    file.seek(0)


    st.download_button(
        "⬇️ Download Example Resume",
        data=file,
        file_name=f"{role}_Sample_Resume.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
